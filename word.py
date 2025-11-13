from docx import Document

# Создаем документ
doc = Document()

# Титульный лист
doc.add_paragraph("\n\n\n\n\n")
doc.add_paragraph("Инновационный колледж АйТи и Бизнеса", style='Title').alignment = 1
doc.add_paragraph("\n\n")
doc.add_paragraph("Реферат", style='Title').alignment = 1
doc.add_paragraph("\n")
doc.add_paragraph("Тема: Влияние спортивной одежды и обуви на эффективность и безопасность занятий физической культурой", style='Subtitle').alignment = 1
doc.add_paragraph("\n\n")
doc.add_paragraph("Выполнила: Батырбекова Алия", style='Normal').alignment = 1
doc.add_paragraph("Преподаватель: Кочкорбаева Н.К.", style='Normal').alignment = 1
doc.add_paragraph("\n\n")
doc.add_paragraph("Город: Манас", style='Normal').alignment = 1
doc.add_paragraph("Дата: 31.10.2025", style='Normal').alignment = 1
doc.add_page_break()

# Содержание (пример)
doc.add_paragraph("Содержание", style='Heading 1')
doc.add_paragraph("Введение..............................................3")
doc.add_paragraph("Глава 1. История спортивной одежды и обуви........4")
doc.add_paragraph("Глава 2. Физиологические требования к одежде.....7")
doc.add_paragraph("Глава 3. Спортивная обувь и безопасность.........10")
doc.add_paragraph("Глава 4. Психологическое влияние..................13")
doc.add_paragraph("Заключение..........................................17")
doc.add_paragraph("Список литературы..................................19")
doc.add_page_break()

# Главы (можно заменить текст на полный реферат)
chapters = [
    "Глава 1. История спортивной одежды и обуви",
    "Глава 2. Физиологические требования к спортивной одежде",
    "Глава 3. Спортивная обувь и безопасность занятий",
    "Глава 4. Психологическое влияние спортивной одежды",
    "Заключение",
    "Список литературы"
]

for chapter in chapters:
    doc.add_paragraph(chapter, style='Heading 1')
    doc.add_paragraph("Текст главы пока будет вставлен здесь...")
    doc.add_page_break()

# Сохраняем документ
doc.save("Реферат_Спортивная_одежда_и_обувь.docx")

print("Файл 'Реферат_Спортивная_одежда_и_обувь.docx' создан!")
